from io import BytesIO

import pymupdf
import pytest
from docx import Document as DocxDocument

from app.parsers.text import DocumentParseError, TextDocumentParser


def test_parse_pdf_keeps_page_markers():
    pdf = pymupdf.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "VPN account application requires manager approval.")

    parsed = TextDocumentParser().parse("vpn.pdf", pdf.tobytes())

    assert parsed.parser == "pymupdf"
    assert parsed.page_count == 1
    assert "--- 第 1 页 ---" in parsed.content
    assert "VPN account application" in parsed.content


def test_parse_docx_includes_paragraphs_and_tables():
    document = DocxDocument()
    document.add_paragraph("VPN 账号申请流程")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "审批人"
    table.rows[0].cells[1].text = "直属主管"
    output = BytesIO()
    document.save(output)

    parsed = TextDocumentParser().parse("vpn.docx", output.getvalue())

    assert parsed.parser == "python-docx"
    assert "VPN 账号申请流程" in parsed.content
    assert "审批人 | 直属主管" in parsed.content


def test_parse_markdown_keeps_headings():
    parsed = TextDocumentParser().parse("vpn.md", "# VPN\n\n账号申请流程".encode())

    assert parsed.parser == "utf-8"
    assert parsed.content == "# VPN\n\n账号申请流程"


def test_scanned_pdf_without_text_is_rejected():
    pdf = pymupdf.open()
    pdf.new_page()

    with pytest.raises(DocumentParseError, match="no extractable text"):
        TextDocumentParser().parse("scan.pdf", pdf.tobytes())
