"""第一版文件文字提取：PDF、DOCX 和 Markdown。"""

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import pymupdf
from docx import Document as DocxDocument


class DocumentParseError(ValueError):
    """上传文件无法提取出可用于入库的文字。"""


@dataclass(frozen=True)
class ParsedDocument:
    """文件解析后的统一文本结果。"""

    content: str
    parser: str
    page_count: int | None = None


class TextDocumentParser:
    """只提取文件本身的文字层，不调用 OCR 或视觉模型。"""

    supported_extensions = {".pdf", ".docx", ".md", ".markdown"}

    def parse(self, filename: str, content: bytes) -> ParsedDocument:
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            return self._parse_pdf(content)
        if suffix == ".docx":
            return self._parse_docx(content)
        if suffix in {".md", ".markdown"}:
            return self._parse_markdown(content)
        raise DocumentParseError(f"unsupported file type: {suffix or 'unknown'}")

    @staticmethod
    def _parse_pdf(content: bytes) -> ParsedDocument:
        try:
            pdf = pymupdf.open(stream=content, filetype="pdf")
        except Exception as error:
            raise DocumentParseError("invalid PDF file") from error

        try:
            page_count = len(pdf)
            pages = []
            for index, page in enumerate(pdf, start=1):
                text = page.get_text("text", sort=True).strip()
                if text:
                    pages.append(f"--- 第 {index} 页 ---\n{text}")
        finally:
            pdf.close()
        return TextDocumentParser._result("\n\n".join(pages), "pymupdf", page_count)

    @staticmethod
    def _parse_docx(content: bytes) -> ParsedDocument:
        try:
            document = DocxDocument(BytesIO(content))
        except Exception as error:
            raise DocumentParseError("invalid DOCX file") from error

        parts = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                line = " | ".join(cell for cell in cells if cell)
                if line:
                    parts.append(line)
        return TextDocumentParser._result("\n\n".join(parts), "python-docx")

    @staticmethod
    def _parse_markdown(content: bytes) -> ParsedDocument:
        try:
            text = content.decode("utf-8-sig")
        except UnicodeDecodeError as error:
            raise DocumentParseError("Markdown must use UTF-8 encoding") from error
        return TextDocumentParser._result(text, "utf-8")

    @staticmethod
    def _result(content: str, parser: str, page_count: int | None = None) -> ParsedDocument:
        normalized = content.strip()
        if not normalized:
            raise DocumentParseError(
                "no extractable text found; scanned PDFs and image text require OCR"
            )
        return ParsedDocument(content=normalized, parser=parser, page_count=page_count)
